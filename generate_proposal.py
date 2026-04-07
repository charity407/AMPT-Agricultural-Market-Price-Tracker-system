"""
AMPT Combined Proposal Generator
Generates COMBINED_PROPOSAL.docx using python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import copy
import os

OUTPUT_PATH = "/home/mukanga/Agri-Price-Tracker/AMPT-Agricultural-Market-Price-Tracker-system/COMBINED_PROPOSAL.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(para, before=0, after=0, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
                     line_val=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_val:
        pf.line_spacing_rule = line_rule
        pf.line_spacing = line_val
    else:
        pf.line_spacing_rule = line_rule


def body_para(doc, text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              keep=False, indent=None):
    """Add a normal body paragraph in Times New Roman 12 with 1.5 spacing."""
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before=0, after=6)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if keep:
        pf.keep_with_next = True
    if indent is not None:
        pf.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def heading1(doc, text):
    """Chapter heading — bold 14pt Times New Roman, centred, spacing above."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=24, after=12)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p


def heading2(doc, text):
    """Section heading — bold 12pt Times New Roman, left-aligned."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=12, after=6)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p


def heading3(doc, text):
    """Sub-section heading — bold italic 12pt Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=8, after=4)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_font(run, size=12, bold=True, italic=True)
    return p


def bullet_para(doc, text, level=0):
    """Bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=0, after=3)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run)
    return p


def numbered_para(doc, text):
    """Numbered list paragraph."""
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=3)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run)
    return p


def code_block(doc, lines):
    """Monospaced Courier New block inside a shaded paragraph."""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, before=0, after=0)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Inches(0.25)
        # Shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return


def diagram_block(doc, lines):
    """ASCII diagram in a bordered Courier New block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=6, after=6)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Inches(0.1)
    pf.right_indent = Inches(0.1)

    # Add border to paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "888888")
        pBdr.append(bdr)
    pPr.append(pBdr)

    # Shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F8F8F8")
    pPr.append(shd)

    full_text = "\n".join(lines)
    run = p.add_run(full_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    return p


def add_page_number(doc):
    """Add page numbers to footer of the current section."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def set_roman_footer(doc, numeral):
    """Put a static roman-numeral text in footer."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(numeral)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_section_break(doc, break_type=WD_SECTION.NEW_PAGE):
    """Add a section break (returns new section)."""
    new_section = doc.add_section(break_type)
    return new_section


def set_margins(section, top=1, bottom=1, left=1, right=1):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def style_table(table, header_bg="1a4731", header_fg="FFFFFF"):
    """Style a table with header row shading."""
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), header_bg)
                tcPr.append(shd)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(
                            *[int(header_fg[i:i+2], 16) for i in (0, 2, 4)]
                        )
                        run.font.bold = True
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_cell(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # Default styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # -----------------------------------------------------------------------
    # SECTION 1: Title Page (roman i)
    # -----------------------------------------------------------------------
    sec1 = doc.sections[0]
    set_margins(sec1)
    set_roman_footer(doc, "i")

    # Institution header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=0, after=4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("CHUKA UNIVERSITY")
    set_font(r, size=14, bold=True)

    p = body_para(doc, "School of Computing and Informatics", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p = body_para(doc, "Department of Computer Science", align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()  # spacer

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=24, after=6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("AGRICULTURAL MARKET PRICE TRACKER SYSTEM (KENYA)")
    set_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=4, after=4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("AgriPrice KE — Project Proposal")
    set_font(r, size=13, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # Course info block
    info_lines = [
        ("Course Code:", "COSC 381"),
        ("Course Title:", "Team Work Project"),
        ("Academic Year:", "2025/2026"),
        ("Submission Date:", "April 2026"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        set_para_spacing(p, before=2, after=2)
        r = p.add_run(f"{label} ")
        set_font(r, bold=True)
        r2 = p.add_run(value)
        set_font(r2)

    doc.add_paragraph()

    # Group Members table
    p = body_para(doc, "GROUP MEMBERS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    members = [
        ("No.", "Adm. Number", "Full Name", "Role"),
        ("1", "EB1/67232/23", "Mukanga Lewis Njega", "Price Management & UI Design"),
        ("2", "EB1/66861/23", "Muigai Charity Wanjiku", "Database Design & Setup"),
        ("3", "EB1/66863/23", "Mwangi Lucy Ann Muthoni", "[Module / Role]"),
        ("4", "EB1/66795/23", "Jason Mageto Paul", "Authentication & User Management"),
        ("5", "EB1/67362/23", "Levi Kisaka Wanyonyi", "Backend Development & Products Module"),
        ("6", "EB1/66902/23", "Oyugi Dominic", "[Module / Role]"),
        ("7", "EB1/66802/23", "Gloria Ndanu", "[Module / Role]"),
    ]

    tbl = doc.add_table(rows=len(members), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(0.4), Inches(1.4), Inches(2.2), Inches(2.5)]
    for i, row_data in enumerate(members):
        row = tbl.rows[i]
        for j, (cell_text, width) in enumerate(zip(row_data, col_widths)):
            row.cells[j].width = width
            table_cell(row.cells[j], cell_text,
                       bold=(i == 0),
                       align=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    style_table(tbl)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_para_spacing(p, before=6, after=2)
    r = p.add_run("Supervisor: ")
    set_font(r, bold=True)
    r2 = p.add_run("[Supervisor Name]")
    set_font(r2)

    # -----------------------------------------------------------------------
    # SECTION 2: Table of Contents (roman ii)
    # -----------------------------------------------------------------------
    add_section_break(doc)
    sec2 = doc.sections[-1]
    set_margins(sec2)
    set_roman_footer(doc, "ii")

    heading1(doc, "TABLE OF CONTENTS")

    toc_entries = [
        ("Abstract", "iii"),
        ("Chapter One: Introduction", "1"),
        ("    1.1 Background of the Study", "1"),
        ("    1.2 Problem Statement", "2"),
        ("    1.3 General Objective", "3"),
        ("    1.4 Specific Objectives", "3"),
        ("    1.5 Research Questions", "4"),
        ("    1.6 Justification", "4"),
        ("    1.7 Proposed Solution", "5"),
        ("Chapter Two: Literature Review", "6"),
        ("    2.1 Introduction", "6"),
        ("    2.2 Review of Similar Systems", "6"),
        ("    2.3 Comparative Analysis", "10"),
        ("    2.4 Gaps Identified and AMPT's Response", "11"),
        ("    2.5 Summary", "12"),
        ("Chapter Three: Methodology", "13"),
        ("    3.1 Research Design", "13"),
        ("    3.2 System Development Methodology", "13"),
        ("    3.3 System Architecture Overview", "14"),
        ("    3.4 Data Flow", "14"),
        ("    3.5 Tools and Technologies", "15"),
        ("    3.6 Project Schedule", "16"),
        ("Chapter Four: Requirements Analysis", "17"),
        ("    4.1 Introduction", "17"),
        ("    4.2 Functional Requirements", "17"),
        ("    4.3 Non-Functional Requirements", "20"),
        ("    4.4 Domain Requirements", "21"),
        ("    4.5 Domain Model", "21"),
        ("    4.6 Use Case Descriptions", "23"),
        ("    4.7 UML Use Case Diagrams", "27"),
        ("Chapter Five: System Design Specification", "30"),
        ("    5.1 System Architecture", "30"),
        ("    5.2 Database Design", "31"),
        ("    5.3 Flowcharts and Activity Diagrams", "34"),
        ("    5.4 Interface Design Overview", "37"),
        ("Chapter Six: System Implementation and Testing", "38"),
        ("    6.1 Implementation Language", "38"),
        ("    6.2 Key Language Features Used", "38"),
        ("    6.3 Search Algorithms Implemented", "39"),
        ("    6.4 Testing Strategy", "44"),
        ("    6.5 Source Code", "45"),
        ("Chapter Seven: Conclusion and Recommendation", "46"),
        ("    7.1 Conclusion", "46"),
        ("    7.2 Recommendations", "47"),
        ("References", "49"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        set_para_spacing(p, before=2, after=2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT)

        r1 = p.add_run(entry)
        set_font(r1, size=11, bold=entry.startswith("Chapter") or entry == "Abstract" or entry == "References")

        r2 = p.add_run(f"\t{page}")
        set_font(r2, size=11)

    # -----------------------------------------------------------------------
    # SECTION 3: Abstract (roman iii)
    # -----------------------------------------------------------------------
    add_section_break(doc)
    sec3 = doc.sections[-1]
    set_margins(sec3)
    set_roman_footer(doc, "iii")

    heading1(doc, "ABSTRACT")

    abstract_text = (
        "Agriculture is the backbone of Kenya's economy, employing over 40% of the total "
        "population and contributing approximately 26% of the Gross Domestic Product (GDP). "
        "Despite this significance, smallholder farmers and market participants continue to "
        "operate with limited access to real-time or near-real-time price information. The "
        "result is a persistent information asymmetry that disadvantages farmers, enables "
        "price manipulation by middlemen, and undermines food security planning by government "
        "agencies."
    )
    body_para(doc, abstract_text)

    abstract_text2 = (
        "This project proposes and presents the Agricultural Market Price Tracker (AMPT) — "
        "a web-based information system, branded AgriPrice KE, designed to monitor, record, "
        "compare, and disseminate agricultural commodity prices across Kenyan markets. The "
        "system is built using Java Servlets, JavaServer Pages (JSP), JDBC, and PostgreSQL, "
        "following the Model-View-Controller (MVC) architectural pattern, hosted on Apache "
        "Tomcat 10."
    )
    body_para(doc, abstract_text2)

    abstract_text3 = (
        "AMPT supports multiple user roles — Administrators, Market Agents, Farmers, Traders, "
        "and Consumers. Market agents record daily commodity prices at their respective markets. "
        "Farmers and other stakeholders can query current prices, compare prices across markets, "
        "track price trends over time, set configurable price alerts, and download summary "
        "reports. Administrators manage the user base, product catalogue, market registry, and "
        "system configuration."
    )
    body_para(doc, abstract_text3)

    abstract_text4 = (
        "The system covers 8 Kenyan regions, 29+ markets, and a catalogue of 35+ agricultural "
        "commodities spanning cereals, legumes, vegetables, fruits, root crops, livestock, dairy, "
        "poultry, and cash crops. The system addresses critical gaps identified in existing "
        "solutions, particularly the lack of a centralised, role-aware, and locally-contextualised "
        "platform for Kenyan agricultural price intelligence."
    )
    body_para(doc, abstract_text4)

    # Keywords
    p = body_para(doc, "")
    r1 = p.add_run("Keywords: ")
    set_font(r1, bold=True)
    r2 = p.add_run(
        "Agricultural market information system, price tracker, Kenya, smallholder farmers, "
        "role-based access control, Java Servlet, PostgreSQL, MVC architecture, price alerts, "
        "price trend analysis."
    )
    set_font(r2, italic=True)

    # -----------------------------------------------------------------------
    # SECTION 4: Chapter One (arabic 1)
    # -----------------------------------------------------------------------
    add_section_break(doc)
    sec4 = doc.sections[-1]
    set_margins(sec4)
    add_page_number(doc)

    # --- CHAPTER ONE ---
    heading1(doc, "CHAPTER ONE: INTRODUCTION")

    heading2(doc, "1.1 Background of the Study")

    body_para(doc,
        "Kenya's agricultural sector plays a central role in the nation's socioeconomic fabric. "
        "The sector supports livelihoods for millions of rural households and drives export "
        "earnings through commodities such as tea, coffee, flowers, and horticultural produce. "
        "However, despite this prominence, access to timely and accurate market price information "
        "remains a critical challenge for the majority of agricultural stakeholders, particularly "
        "smallholder farmers."
    )
    body_para(doc,
        "Market information systems have been studied extensively as tools for improving market "
        "efficiency. The fundamental premise is straightforward: when farmers, traders, and "
        "consumers have access to the same price information, markets function more competitively, "
        "middlemen's price arbitrage is reduced, and farmers are empowered to make better-informed "
        "selling decisions. Mobile technology and internet penetration in Kenya, now exceeding "
        "22 million internet users (CA Kenya, 2024), creates a viable infrastructure for digital "
        "market information delivery."
    )
    body_para(doc,
        "Historically, agricultural price dissemination in Kenya relied on radio broadcasts, "
        "printed market bulletins, and informal word-of-mouth networks. These channels suffer from "
        "significant time lags, geographic limitations, and lack of granularity. The Kenya Markets "
        "Trust, the East Africa Grain Council, and the National Farmers' Information Service "
        "(NAFIS) have made efforts to provide price data, but these remain fragmented, "
        "infrequently updated, and not easily accessible to end users in real time."
    )
    body_para(doc,
        "The emergence of web-based systems and mobile applications has created new opportunities "
        "to bridge this information gap. This project responds to that opportunity by developing "
        "AMPT — a comprehensive, web-based price monitoring and dissemination system tailored to "
        "the Kenyan agricultural context."
    )

    heading2(doc, "1.2 Problem Statement")

    body_para(doc,
        "Smallholder farmers and other agricultural market participants in Kenya consistently face "
        "an information asymmetry regarding commodity prices. Farmers often sell produce at markets "
        "without knowing prices at competing markets, while buyers exploit this ignorance to "
        "negotiate unfairly low prices. The consequences include:"
    )
    bullet_para(doc, "Income losses for farmers who sell at below-market prices due to lack of comparative price knowledge.")
    bullet_para(doc, "Inefficient market allocation where produce travels long distances unnecessarily because price signals are absent or distorted.")
    bullet_para(doc, "Food security risks arising from poor production planning driven by price uncertainty.")
    bullet_para(doc, "Lack of accountability in market pricing, enabling cartel behaviour and price fixing.")
    bullet_para(doc, "Inadequate policy data for government agencies that need timely price trends to make subsidy, intervention, and food importation decisions.")
    body_para(doc,
        "Existing digital solutions either focus on a narrow set of commodities, operate at a "
        "national aggregate level without market-level granularity, lack role-based access for "
        "data entry integrity, or are not maintained with current data. There is therefore a need "
        "for a dedicated, locally-contextualised, web-based system that collects, stores, and "
        "presents agricultural market price data at the individual market level, across the full "
        "breadth of Kenyan agricultural commodities, with appropriate access controls and "
        "analytical features."
    )

    heading2(doc, "1.3 General Objective")
    body_para(doc,
        "To design and implement a web-based Agricultural Market Price Tracker (AMPT) system that "
        "enables real-time collection, storage, comparison, and dissemination of agricultural "
        "commodity prices across Kenyan markets, thereby reducing information asymmetry among "
        "farmers, traders, and other market participants."
    )

    heading2(doc, "1.4 Specific Objectives")
    numbered_para(doc,
        "To analyse the existing information needs of agricultural market stakeholders in Kenya "
        "and identify functional requirements for a price tracking system."
    )
    numbered_para(doc,
        "To design a relational database schema capable of storing price data, market metadata, "
        "commodity catalogues, user accounts, and alert configurations for the Kenyan agricultural "
        "market context."
    )
    numbered_para(doc,
        "To implement a role-based web application using Java Servlets and JSP that supports "
        "price entry by market agents, price querying by farmers and consumers, and administration "
        "of the system by authorised administrators."
    )
    numbered_para(doc,
        "To implement price comparison, trend analysis, alert notification, and reporting features "
        "that enable data-driven decision making for agricultural market participants."
    )
    numbered_para(doc,
        "To implement and demonstrate standard search and sorting algorithms within the system to "
        "facilitate efficient retrieval and presentation of price data."
    )
    numbered_para(doc,
        "To test the system against defined functional and non-functional requirements to validate "
        "correctness, usability, and performance."
    )

    heading2(doc, "1.5 Research Questions")
    numbered_para(doc,
        "What are the primary information needs of smallholder farmers and market agents regarding "
        "agricultural commodity pricing in Kenya, and what features must a price tracking system "
        "possess to address those needs effectively?"
    )
    numbered_para(doc,
        "How can a role-based web system be designed and implemented to ensure that price data "
        "entered by market agents is accurate, consistent, and free from unauthorised manipulation?"
    )
    numbered_para(doc,
        "What database design and query optimisation strategies are most appropriate for storing "
        "and retrieving large volumes of time-series agricultural price data across multiple "
        "markets and commodities?"
    )
    numbered_para(doc,
        "How effective are configurable price alert mechanisms in informing farmers of significant "
        "price movements, and what threshold parameters yield the most actionable notifications?"
    )
    numbered_para(doc,
        "To what extent does access to historical price trend data and multi-market price "
        "comparisons influence the selling and buying decisions of agricultural market participants?"
    )
    numbered_para(doc,
        "What algorithms and system architecture choices best balance response time performance "
        "and scalability for a multi-user agricultural price tracking web application?"
    )

    heading2(doc, "1.6 Justification")

    heading3(doc, "Economic Impact")
    body_para(doc,
        "Price information asymmetry costs Kenyan smallholder farmers an estimated 10–30% of "
        "potential income per season (World Bank, 2022). A system that reduces this asymmetry "
        "directly translates to improved farm incomes and better living standards for rural "
        "households."
    )

    heading3(doc, "Policy Relevance")
    body_para(doc,
        "Kenya's Big Four Agenda and the subsequent Bottom-Up Economic Transformation Agenda "
        "(BETA) identify food security as a pillar. Accurate, real-time price data enables "
        "evidence-based policy interventions, subsidies, and strategic food reserves management."
    )

    heading3(doc, "Academic Contribution")
    body_para(doc,
        "This project demonstrates the application of core software engineering principles — "
        "requirements analysis, object-oriented design, MVC architecture, relational database "
        "design, and algorithm implementation — to a real-world, socially significant problem "
        "domain."
    )

    heading3(doc, "Technological Feasibility")
    body_para(doc,
        "Kenya's growing internet penetration (>50% urban, >20% rural) and the proliferation of "
        "Android smartphones at affordable price points make web-based market information systems "
        "practically viable for the target users."
    )

    heading3(doc, "Institutional Relevance")
    body_para(doc,
        "Chuka University is located in a primarily agricultural region (Tharaka-Nithi County). "
        "The institution and its surrounding community are direct potential beneficiaries of such "
        "a system, giving this project immediate local relevance beyond its academic value."
    )

    heading2(doc, "1.7 Proposed Solution")
    body_para(doc,
        "The proposed solution is AgriPrice KE — a web-based Agricultural Market Price Tracker "
        "built on the following design principles:"
    )
    bullet_para(doc,
        "Role-Based Access: Five distinct user roles (Admin, Market Agent, Farmer, Trader, "
        "Consumer) with permissions tailored to each role's data access and entry needs."
    )
    bullet_para(doc,
        "Market-Level Granularity: Prices are recorded and queried at the individual market "
        "level, linked to a geographic region hierarchy (Region → Market)."
    )
    bullet_para(doc,
        "Comprehensive Commodity Catalogue: 35+ commodities across 9 categories, covering the "
        "full breadth of Kenyan agricultural production."
    )
    bullet_para(doc,
        "Analytical Features: Multi-market price comparison, time-series trend analysis, "
        "configurable price alerts, and downloadable reports."
    )
    bullet_para(doc,
        "Standard Technology Stack: Java 17 Servlets + JSP (MVC), JDBC, PostgreSQL 16, Apache "
        "Tomcat 10 — no heavyweight frameworks — demonstrating foundational software engineering "
        "competencies."
    )
    bullet_para(doc,
        "Algorithm Demonstration: Linear search, binary search, bubble sort, and quicksort "
        "implemented in Java within the application logic, with complexity analysis documented."
    )

    # -----------------------------------------------------------------------
    # CHAPTER TWO
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER TWO: LITERATURE REVIEW")

    heading2(doc, "2.1 Introduction")
    body_para(doc,
        "This chapter reviews existing agricultural market information systems — both globally "
        "and within the East African context — with particular focus on their features, "
        "limitations, and the gaps that motivate the development of AMPT. The review draws on "
        "deployed systems, academic research, and reports from agricultural development "
        "organisations. For each system, this chapter highlights both notable strengths that "
        "informed AMPT's design and limitations that create the need for the system."
    )

    heading2(doc, "2.2 Review of Similar Systems")

    heading3(doc, "2.2.1 Agricultural Commodity Exchange (ACE) — Kenya")
    body_para(doc,
        "The Agricultural Commodity Exchange is a Kenya-based electronic trading platform "
        "established to provide market information and facilitate agricultural commodity trading."
    )
    p = body_para(doc, "")
    r = p.add_run("Strengths (Informing AMPT's Design):")
    set_font(r, bold=True)
    bullet_para(doc,
        "ACE pioneered Kenya's first electronic trading platform for agricultural commodities, "
        "demonstrating the viability of digital market systems in the Kenyan context — a proof "
        "of concept AMPT builds upon."
    )
    bullet_para(doc,
        "Its warehouse receipt system established a trusted, structured approach to commodity "
        "verification that inspired AMPT's data validation service design."
    )
    bullet_para(doc,
        "ACE's direct connection between farmers and buyers showed that digital platforms can "
        "successfully bridge the producer–buyer gap in Kenya's agricultural sector."
    )
    p = body_para(doc, "")
    r = p.add_run("Limitations Addressed by AMPT:")
    set_font(r, bold=True)
    bullet_para(doc,
        "Limited to formally registered farmers and traders, excluding the majority of Kenya's "
        "smallholder farming community."
    )
    bullet_para(doc, "Requires significant capital investment, making it inaccessible to low-income farmers.")
    bullet_para(doc, "Focuses primarily on trading rather than broad price information dissemination.")
    bullet_para(doc, "Does not provide comprehensive trend analysis tools or historical comparisons.")

    heading3(doc, "2.2.2 NAFIS — National Farmers' Information Service (Kenya)")
    body_para(doc,
        "A Kenyan government portal (nafis.go.ke) managed by the Ministry of Agriculture, "
        "providing extension information and some market pricing data."
    )
    p = body_para(doc, "")
    r = p.add_run("Strengths (Informing AMPT's Design):")
    set_font(r, bold=True)
    bullet_para(doc,
        "As a government-backed initiative, NAFIS established the principle of free public "
        "access to agricultural price data — a core value AMPT adopts and expands."
    )
    bullet_para(doc,
        "Its multi-commodity price reporting across selected Kenyan markets provided a "
        "foundational data model improved with real-time updates and wider coverage."
    )
    p = body_para(doc, "")
    r = p.add_run("Limitations Addressed by AMPT:")
    set_font(r, bold=True)
    bullet_para(doc, "Price data is irregularly updated and often weeks behind current market conditions.")
    bullet_para(doc, "No role-based access — no mechanism for market agents to directly input prices.")
    bullet_para(doc, "No price trend visualisation or historical data exploration.")
    bullet_para(doc, "No price alert or notification system.")
    bullet_para(doc, "Interface is not optimised for mobile devices.")

    heading3(doc, "2.2.3 East Africa Grain Council (EAGC) — Price Monitoring System")
    body_para(doc,
        "The EAGC operates a grain price monitoring network across East Africa, publishing "
        "weekly price bulletins for maize, wheat, beans, and related commodities."
    )
    p = body_para(doc, "")
    r = p.add_run("Strengths: ")
    set_font(r, bold=True)
    r2 = p.add_run(
        "Regional price comparisons and seasonal price trend reports demonstrated the value "
        "of cross-market data aggregation."
    )
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Limitations Addressed by AMPT:")
    set_font(r, bold=True)
    bullet_para(doc, "Restricted to grain commodities — does not cover vegetables, fruits, livestock, dairy, or poultry.")
    bullet_para(doc, "Weekly publication cycle is too infrequent for daily market decisions.")
    bullet_para(doc, "No user accounts, personalised alerts, or role-based data entry.")
    bullet_para(doc, "Primarily designed for policy analysts, not individual farmers.")

    heading3(doc, "2.2.4 Esoko — Regional Platform (West/East Africa)")
    body_para(doc,
        "A mobile-based agricultural information platform operating in West Africa and limited "
        "East African countries, providing price information via SMS and web."
    )
    p = body_para(doc, "")
    r = p.add_run("Strengths (Informing AMPT's Design):")
    set_font(r, bold=True)
    bullet_para(doc,
        "Esoko proved that mobile delivery of market prices can reach rural farming communities "
        "lacking desktop internet access — validating AMPT's mobile-first design strategy."
    )
    bullet_para(doc,
        "Its multi-source market data aggregation approach directly informs AMPT's Kenya Market "
        "Data Aggregator component."
    )
    p = body_para(doc, "")
    r = p.add_run("Limitations Addressed by AMPT:")
    set_font(r, bold=True)
    bullet_para(doc,
        "Commercial SaaS model with per-user licensing costs — inaccessible for small farmer "
        "cooperatives or county government deployment."
    )
    bullet_para(doc,
        "SMS-centric design is outdated given Kenya's smartphone penetration exceeding 50%."
    )
    bullet_para(doc, "Limited historical data and no trend analysis tools.")
    bullet_para(doc, "Dependent on external data providers with limited Kenya-specific market coverage.")

    heading3(doc, "2.2.5 E-NAM (Electronic National Agriculture Market) — India")
    body_para(doc,
        "India's unified online agriculture market platform covering multiple states, connecting "
        "farmers, traders, and buyers through a centralised digital marketplace."
    )
    p = body_para(doc, "")
    r = p.add_run("Strengths (Informing AMPT's Design):")
    set_font(r, bold=True)
    bullet_para(doc,
        "E-NAM demonstrated at scale how technology can unify fragmented agricultural markets "
        "— a key lesson for AMPT's multi-market integration approach."
    )
    bullet_para(doc,
        "Its GIS integration for spatial price visualisation directly inspired AMPT's regional "
        "comparison features."
    )
    bullet_para(doc,
        "E-NAM's real-time bidding and price transparency model validated the core premise of "
        "AMPT: eliminating information asymmetry improves farmer income."
    )
    p = body_para(doc, "")
    r = p.add_run("Limitations for the Kenyan Context:")
    set_font(r, bold=True)
    bullet_para(doc,
        "Requires mandatory market participation structures (APMC mandis) that do not exist in "
        "Kenya's informal market ecosystem."
    )
    bullet_para(doc, "Not designed for Kenya's specific crops, livestock categories, and seasonal patterns.")
    bullet_para(doc, "Assumes high digital infrastructure not reflecting Kenya's rural connectivity realities.")

    heading3(doc, "2.2.6 Reuters Market Light (India) and Farmcrowdy (Nigeria)")
    body_para(doc,
        "Reuters Market Light (2007–2015) provided personalised SMS price alerts and crop "
        "advisories to Indian farmers. Though successful in concept, it was discontinued in "
        "2015 due to commercial sustainability challenges — a lesson for AMPT's open, "
        "institutionally deployable model. Farmcrowdy Price Tracker in Nigeria offers price "
        "listings and trend data but is Nigeria-specific and lacks cross-market comparison and "
        "configurable alerts."
    )

    heading3(doc, "2.2.7 Academic Prototypes")
    body_para(doc,
        "Several academic studies have proposed agricultural market information systems. "
        "Ochieng et al. (2020) developed a mobile-based price notification system for maize "
        "farmers in Western Kenya using Android and MySQL. Mutua & Mugo (2021) proposed a "
        "web-based system using PHP/Laravel for horticultural price tracking in Central Kenya."
    )
    body_para(doc,
        "Common limitations: narrow commodity or geographic scope, no role hierarchy, no trend "
        "analysis or alert features, and no deployment beyond the academic setting."
    )

    heading2(doc, "2.3 Comparative Analysis")
    body_para(doc,
        "The table below summarises how existing systems compare to the proposed AMPT system:"
    )

    # Comparison table
    comp_headers = ["Feature", "ACE", "NAFIS", "EAGC", "Esoko", "E-NAM", "AMPT"]
    comp_rows = [
        ["Real-time Prices", "Limited", "No", "No", "Yes", "Yes", "Yes"],
        ["Kenya Market Focus", "Yes", "Yes", "Partial", "Limited", "No", "Yes"],
        ["Smallholder Accessibility", "No", "Limited", "No", "Limited", "No", "High"],
        ["Comprehensive Commodity Coverage", "No", "Limited", "Grains Only", "Limited", "Partial", "Yes (35+)"],
        ["Role-Based Data Entry", "No", "No", "No", "No", "No", "Yes"],
        ["Interactive Visualisation", "No", "No", "No", "Limited", "No", "Advanced"],
        ["Configurable Price Alerts", "No", "No", "No", "Limited", "No", "Yes"],
        ["Free/Open Deployment", "No", "Yes", "Yes", "No", "Yes", "Yes"],
        ["Mobile Responsive", "Limited", "No", "No", "Yes", "Limited", "Optimised"],
        ["Seasonal Analysis", "No", "No", "No", "No", "No", "Yes"],
    ]
    all_rows = [comp_headers] + comp_rows
    tbl2 = doc.add_table(rows=len(all_rows), cols=len(comp_headers))
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(1.6), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.7)]
    for i, row_data in enumerate(all_rows):
        row = tbl2.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].width = col_w[j] if j < len(col_w) else Inches(0.7)
            is_header = (i == 0)
            is_ampt = (j == len(comp_headers) - 1) and not is_header
            table_cell(row.cells[j], val, bold=(is_header or is_ampt), size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    style_table(tbl2)

    heading2(doc, "2.4 Gaps Identified and AMPT's Response")
    numbered_para(doc,
        "Lack of Real-Time Multi-Market Integration: No existing system combines real-time, "
        "publicly accessible, multi-market price data covering both formal and informal Kenyan "
        "markets. AMPT addresses this through daily agent-submitted price entries across 29+ "
        "markets."
    )
    numbered_para(doc,
        "Inaccessibility for Smallholder Farmers: No existing system combines free access, "
        "mobile-first design, and intuitive interface for users with varying digital literacy. "
        "AMPT's low-bandwidth, mobile-first design directly addresses this."
    )
    numbered_para(doc,
        "Narrow Commodity Coverage: No single Kenyan system comprehensively tracks cereals, "
        "pulses, horticulture, tea, coffee, livestock, and dairy in one platform. AMPT covers "
        "35+ commodities across 9 EAC-aligned categories."
    )
    numbered_para(doc,
        "Absence of Role-Based Data Integrity: Existing systems either have no data entry "
        "mechanism or no access controls. AMPT enforces a 5-role hierarchy ensuring only "
        "verified Market Agents submit price data."
    )
    numbered_para(doc,
        "No Configurable User Alerts: AMPT allows any user to set threshold-based alerts "
        "(ABOVE/BELOW) for any product-market combination, a feature absent from all reviewed "
        "systems."
    )
    numbered_para(doc,
        "Open, Institutionally Deployable Architecture: Unlike proprietary systems, AMPT is "
        "built on open-source technologies with clear documentation, enabling local government "
        "bodies, cooperatives, and universities to maintain and extend it independently."
    )

    heading2(doc, "2.5 Summary")
    body_para(doc,
        "The literature review confirms that while various agricultural market information "
        "systems exist globally and in Africa, none fully addresses the combination of real-time "
        "data entry by verified agents, comprehensive commodity coverage, multi-market "
        "comparison, configurable alerts, and open deployability within the Kenyan context. "
        "AMPT is positioned to bridge these gaps."
    )

    # -----------------------------------------------------------------------
    # CHAPTER THREE
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER THREE: METHODOLOGY")

    heading2(doc, "3.1 Research Design")
    body_para(doc,
        "This project follows an applied research design, combining primary data collection "
        "with secondary research to inform system requirements. The research is both qualitative "
        "(identifying user needs and workflows) and quantitative (evaluating system performance "
        "against defined metrics)."
    )
    p = body_para(doc, "Primary data collection methods include:")
    bullet_para(doc,
        "Semi-structured interviews with market agents, smallholder farmers, and agricultural "
        "extension officers to identify functional needs and workflows."
    )
    bullet_para(doc,
        "Direct observation of price data collection workflows at Thika and Muranga markets "
        "(accessible from Chuka University catchment area)."
    )
    body_para(doc,
        "Secondary research includes review of NAFIS bulletins, EAGC price reports, county "
        "agriculture department records, and academic literature on agricultural market "
        "information systems in sub-Saharan Africa."
    )

    heading2(doc, "3.2 System Development Methodology")
    body_para(doc,
        "AMPT is developed using the Iterative and Incremental Development (IID) model, "
        "structured into four major phases aligned with the project timeline:"
    )
    bullet_para(doc,
        "Phase 1 — Requirements Analysis (Weeks 4–6): Elicit, document, and validate functional "
        "and non-functional requirements using interviews, use case modelling, and domain "
        "modelling."
    )
    bullet_para(doc,
        "Phase 2 — System Design (Weeks 7–8): Produce database schema (ERD), system architecture "
        "diagrams, MVC component design, interface wireframes, and activity diagrams for core "
        "workflows."
    )
    bullet_para(doc,
        "Phase 3 — Implementation and Testing (Weeks 9–12): Iteratively implement system modules "
        "in Java (Servlets + JDBC), with each module tested before integration. Unit testing, "
        "integration testing, and user acceptance testing (UAT)."
    )
    bullet_para(doc,
        "Phase 4 — Presentation and Finalisation (Week 13): System demonstration, report "
        "finalisation, and code submission on CD."
    )
    body_para(doc,
        "Rationale for IID: Agricultural market systems involve evolving requirements as "
        "stakeholder feedback is incorporated. IID allows for mid-development corrections without "
        "the rigidity of the Waterfall model, while providing the structured deliverables "
        "required for academic assessment."
    )

    heading2(doc, "3.3 System Architecture Overview")
    body_para(doc,
        "The AMPT system follows a three-tier architecture implementing the MVC pattern:"
    )
    bullet_para(doc,
        "Tier 1 — Presentation (Frontend): Displays prices, charts, search, and alerts. Built "
        "with Vanilla HTML5/CSS3/JavaScript with PWA offline support — works on basic Android "
        "phones on 2G networks."
    )
    bullet_para(doc,
        "Tier 2 — Application (Backend): Handles all business logic — data validation, user "
        "authentication, aggregation, alert triggering. Pure Java Servlets with no frameworks "
        "— deployable by Kenyan universities and county offices."
    )
    bullet_para(doc,
        "Tier 3 — Data (Database): Stores all price records, market data, user accounts, and "
        "seasonal calendars. PostgreSQL with time-series indexing for fast historical queries."
    )

    heading2(doc, "3.4 Data Flow")
    body_para(doc, "The end-to-end data flow from market to farmer:")
    numbered_para(doc,
        "A verified Market Agent submits the day's price for a commodity at a registered market."
    )
    numbered_para(doc,
        "The backend's Data Validation Service checks the entry against market-specific rules "
        "(flags prices deviating more than 40% from the weekly average as anomalies)."
    )
    numbered_para(doc,
        "The validated entry is written to the PostgreSQL price_entries table with timestamp, "
        "market ID, product ID, and agent ID."
    )
    numbered_para(doc,
        "The Price Alert Service checks whether the new price crosses any user-configured "
        "threshold and dispatches notifications accordingly."
    )
    numbered_para(doc,
        "The most recent prices are cached for offline access by users with intermittent "
        "connectivity."
    )

    heading2(doc, "3.5 Tools and Technologies")

    heading3(doc, "3.5.1 Development Tools")

    dev_tools = [
        ("Tool", "Category", "Usage"),
        ("VS Code / IntelliJ IDEA", "IDE", "Frontend HTML/CSS/JS (VS Code); Java backend development (IntelliJ IDEA)"),
        ("Git & GitHub", "Version Control", "Source code management; team collaboration across feature branches"),
        ("Postman", "API Testing", "Testing REST API endpoints; shareable test collections"),
        ("pgAdmin", "Database GUI", "Designing tables, running SQL queries, verifying indexes"),
        ("JUnit", "Unit Testing", "Automated unit tests for Java backend service classes"),
        ("draw.io / StarUML", "Diagramming", "UML use case diagrams, domain class diagrams, architecture diagrams"),
    ]
    tbl3 = doc.add_table(rows=len(dev_tools), cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w3 = [Inches(1.8), Inches(1.4), Inches(3.3)]
    for i, row_data in enumerate(dev_tools):
        row = tbl3.rows[i]
        for j, val in enumerate(row_data):
            row.cells[j].width = col_w3[j]
            table_cell(row.cells[j], val, bold=(i == 0), size=10)
    style_table(tbl3)
    doc.add_paragraph()

    heading3(doc, "3.5.2 Technology Stack")

    p = body_para(doc, "")
    r = p.add_run("Frontend Technologies:")
    set_font(r, bold=True)
    bullet_para(doc, "HTML5: Structural markup with semantic elements for accessibility.")
    bullet_para(doc, "CSS3: Responsive layout using media queries — adapts from 320px mobile to desktop.")
    bullet_para(doc, "JavaScript (ES6+): Client-side interactivity, REST API calls, offline cache management.")
    bullet_para(doc, "Chart.js: Interactive line and bar charts for price visualisation.")
    bullet_para(doc, "Service Workers & IndexedDB: PWA offline caching of recent price data.")

    p = body_para(doc, "")
    r = p.add_run("Backend Technologies:")
    set_font(r, bold=True)
    bullet_para(doc, "Java SE 17: Backend programming language — strongly typed, stable, widely taught at Kenyan universities.")
    bullet_para(doc, "Jakarta Servlet API 6.0: HTTP request-response handling, session management, MVC controller layer.")
    bullet_para(doc, "JavaServer Pages (JSP) + JSTL: Server-side view rendering.")
    bullet_para(doc, "JDBC (Direct, no ORM): Database access with PreparedStatement — prevents SQL injection.")

    p = body_para(doc, "")
    r = p.add_run("Database Technology:")
    set_font(r, bold=True)
    bullet_para(doc,
        "PostgreSQL 16: Primary relational database. Superior support for time-series indexing "
        "critical for historical price queries. All interactions use parameterised queries."
    )

    p = body_para(doc, "")
    r = p.add_run("Architectural Pattern — MVC:")
    set_font(r, bold=True)
    bullet_para(doc,
        "Model: Java POJO classes (Product, Market, PriceEntry, PriceAlert, Region, Category) "
        "and utility classes (DBConnection, AlgorithmUtils)."
    )
    bullet_para(doc, "View: JSP pages with JSTL tags rendering HTML responses from model attributes.")
    bullet_para(doc,
        "Controller: Java Servlets handling HTTP GET and POST requests, coordinating JDBC "
        "queries and forwarding to JSP views."
    )

    p = body_para(doc, "")
    r = p.add_run("Security Approach:")
    set_font(r, bold=True)
    bullet_para(doc, "SQL Injection prevention via PreparedStatement parameter binding.")
    bullet_para(doc, "Session-based authentication (HttpSession) with role attribute enforcement.")
    bullet_para(doc, "Role-based route protection (admin-only servlets check userRole session attribute).")
    bullet_para(doc, "Soft deletes (is_active flag) to preserve referential integrity and audit trails.")
    bullet_para(doc, "Passwords stored as bcrypt hashes — never in plaintext.")

    heading2(doc, "3.6 Project Schedule")

    sched_rows = [
        ("Weeks", "Activity", "Deliverable"),
        ("1–3", "Proposal writing, technology research", "This proposal document"),
        ("4–6", "Requirements analysis (OOA, use cases, domain model)", "Chapter Four"),
        ("7–8", "System design (architecture, ERD, flowcharts, UML diagrams)", "Chapter Five"),
        ("9–12", "Implementation: Auth, Price CRUD, Alerts, Reports, Testing", "Working system + Chapter Six"),
        ("13", "Final presentation and marking", "Presentation + CD submission"),
    ]
    tbl_s = doc.add_table(rows=len(sched_rows), cols=3)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.LEFT
    sw = [Inches(0.8), Inches(3.0), Inches(2.7)]
    for i, rd in enumerate(sched_rows):
        row = tbl_s.rows[i]
        for j, val in enumerate(rd):
            row.cells[j].width = sw[j]
            table_cell(row.cells[j], val, bold=(i == 0), size=10)
    style_table(tbl_s)

    # -----------------------------------------------------------------------
    # CHAPTER FOUR
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER FOUR: REQUIREMENTS ANALYSIS")

    heading2(doc, "4.1 Introduction")
    body_para(doc,
        "Requirements analysis was conducted using Object-Oriented Analysis (OOA) techniques, "
        "including use case modelling, domain modelling, and stakeholder interviews. This chapter "
        "presents the system's functional, non-functional, and domain requirements, followed by "
        "a domain model and use case specifications with UML use case diagrams."
    )

    heading2(doc, "4.2 Functional Requirements")

    heading3(doc, "4.2.1 User Registration and Authentication")
    req_block = [
        ("FR1", "The system shall allow new users to self-register with a full name, email address, password, phone number, and role selection."),
        ("FR2", "The system shall authenticate registered users via email and password."),
        ("FR3", "Passwords shall be stored as hashed values (bcrypt) — never in plaintext."),
        ("FR4", "The system shall maintain user sessions and enforce login for all protected pages."),
        ("FR5", "The system shall provide a logout mechanism that invalidates the session."),
    ]
    for code, text in req_block:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.2 Role-Based Access Control")
    reqs = [
        ("FR6", "The system shall enforce five user roles: Admin, Market Agent, Farmer, Trader, Consumer."),
        ("FR7", "Admin users shall have access to all modules."),
        ("FR8", "Market Agents shall be able to enter and edit price data for their assigned markets."),
        ("FR9", "Farmers, Traders, and Consumers shall have read-only access to price data."),
        ("FR10", "The system shall redirect unauthorised access attempts to the login page."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.3 Price Data Collection")
    reqs = [
        ("FR11", "Market Agents shall be able to enter a price record specifying: product, market, unit price (KES), date, and optional notes."),
        ("FR12", "The system shall prevent duplicate price entries for the same product-market-date combination using a database unique constraint (UPSERT logic)."),
        ("FR13", "The system shall display a confirmation or error message after each submission."),
        ("FR14", "The system shall flag submitted prices that deviate more than 40% from the weekly average as anomalies for review."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.4 Price Listing and Search")
    reqs = [
        ("FR15", "The system shall display all price entries in a tabular format with columns: Product, Unit Price, Market, Region, and Date."),
        ("FR16", "The system shall support filtering by crop name, market name, region, and date range."),
        ("FR17", "Filters shall use case-insensitive matching."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.5 Multi-Market Price Comparison")
    p = body_para(doc, "")
    r = p.add_run("FR18: ")
    set_font(r, bold=True)
    r2 = p.add_run(
        "The system shall allow users to compare the price of a single commodity across "
        "multiple selected markets simultaneously, displaying the most recent recorded price "
        "for each market."
    )
    set_font(r2)

    heading3(doc, "4.2.6 Price Trend Analysis")
    reqs = [
        ("FR19", "The system shall display a time-series line chart of price history for a selected product-market combination."),
        ("FR20", "The user shall be able to select the time window: 30 days, 90 days, or 180 days."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.7 Price Alerts")
    reqs = [
        ("FR21", "Users shall be able to create a price alert by specifying: product, market (optional), alert direction (ABOVE/BELOW), and a threshold price."),
        ("FR22", "The system shall evaluate active alerts against current prices and flag triggered alerts."),
        ("FR23", "Users shall be able to view, edit, activate/deactivate, and delete their own alerts."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.8 Reporting")
    reqs = [
        ("FR24", "The system shall generate summary reports showing aggregate statistics (count, average, minimum, and maximum price) grouped by product and market."),
        ("FR25", "The system shall support CSV export of price data for download."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.9 Product, Market, and User Administration")
    reqs = [
        ("FR26", "Admin users shall be able to add, view, and soft-delete products, markets, and categories."),
        ("FR27", "Admin users shall be able to view all registered users and update their account status (Active/Suspended) and role."),
    ]
    for code, text in reqs:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading3(doc, "4.2.10 Audit Logging")
    p = body_para(doc, "")
    r = p.add_run("FR28: ")
    set_font(r, bold=True)
    r2 = p.add_run(
        "The system shall record significant data changes (inserts, updates, deletes) in an "
        "audit log table with actor identity, action type, entity affected, and timestamp."
    )
    set_font(r2)

    heading2(doc, "4.3 Non-Functional Requirements")

    nfr_sections = [
        ("Performance", [
            ("NFR1", "The system shall respond to standard page requests within 3 seconds under normal load (up to 50 concurrent users on a local server)."),
            ("NFR2", "Database queries for price listing shall use indexed columns to achieve sub-second query execution."),
        ]),
        ("Security", [
            ("NFR3", "All database interactions shall use PreparedStatement to prevent SQL injection."),
            ("NFR4", "User passwords shall not be stored in plaintext."),
            ("NFR5", "All application pages except login and registration shall require an active authenticated session."),
        ]),
        ("Usability", [
            ("NFR6", "The system interface shall be responsive and usable on both desktop and mobile (smartphone) browsers."),
            ("NFR7", "Form fields shall provide clear labels, placeholders, and validation feedback."),
            ("NFR8", "Navigation shall be consistent across all pages via a persistent header navbar."),
        ]),
        ("Reliability", [
            ("NFR9", "The system shall use try-with-resources for all database connections to prevent connection leaks."),
            ("NFR10", "The system shall gracefully handle database unavailability with a user-friendly error message rather than a stack trace."),
        ]),
        ("Maintainability", [
            ("NFR11", "The system shall follow the MVC pattern to separate business logic, data access, and presentation concerns."),
            ("NFR12", "All database credentials shall be stored in an external properties file (db.properties) not hardcoded in source code."),
        ]),
        ("Scalability", [
            ("NFR13", "The database schema shall use normalised relational tables to support extension to additional regions, commodities, and markets without schema changes."),
        ]),
    ]

    for category, items in nfr_sections:
        p = body_para(doc, "")
        r = p.add_run(category + ":")
        set_font(r, bold=True)
        for code, text in items:
            p2 = body_para(doc, "", indent=0.25)
            r1 = p2.add_run(f"{code}: ")
            set_font(r1, bold=True)
            r2 = p2.add_run(text)
            set_font(r2)

    heading2(doc, "4.4 Domain Requirements")
    dr_items = [
        ("DR1", "Currency — All prices shall be recorded in Kenyan Shillings (KES)."),
        ("DR2", "Units — Each product shall have a standard unit (e.g., \"90kg bag\" for maize, \"kg\" for vegetables, \"litre\" for milk) consistent with Kenyan market conventions."),
        ("DR3", "Date Format — Price dates shall follow ISO 8601 (YYYY-MM-DD)."),
        ("DR4", "Market Status — Only markets with status = 'ACTIVE' shall appear in price entry dropdowns."),
        ("DR5", "Commodity Categories — Products shall be classified into one of nine EAC-aligned categories: Cereals & Grains, Legumes & Pulses, Vegetables, Fruits, Root Crops & Tubers, Livestock, Dairy, Poultry, and Cash Crops."),
    ]
    for code, text in dr_items:
        p = body_para(doc, "")
        r = p.add_run(f"{code}: ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)

    heading2(doc, "4.5 Domain Model")
    body_para(doc,
        "The core domain entities and their relationships are defined below using UML Class "
        "Diagram notation:"
    )

    domain_model_lines = [
        "REGION",
        "  - region_id : int {PK}",
        "  - region_name : String",
        "  - region_code : String",
        "  - country : String",
        "  [1] -----< [many] MARKET",
        "",
        "MARKET",
        "  - market_id : int {PK}",
        "  - market_name : String",
        "  - town : String",
        "  - operating_days : String",
        "  - operating_hours : String",
        "  - status : enum {ACTIVE, INACTIVE}",
        "  - region_id : int {FK -> REGION}",
        "",
        "PRODUCT_CATEGORY",
        "  - category_id : int {PK}",
        "  - category_name : String",
        "  - category_type : enum {CROP, LIVESTOCK, CASH}",
        "  [1] -----< [many] PRODUCT",
        "",
        "PRODUCT",
        "  - product_id : int {PK}",
        "  - product_name : String",
        "  - local_name : String",
        "  - standard_unit : String",
        "  - is_active : boolean",
        "  - category_id : int {FK -> PRODUCT_CATEGORY}",
        "",
        "USER",
        "  - user_id : int {PK}",
        "  - full_name : String",
        "  - email_address : String {UNIQUE}",
        "  - password_hash : String",
        "  - role : enum {ADMIN, AGENT, FARMER, TRADER, CONSUMER}",
        "  - phone_number : String",
        "  - account_status : enum {ACTIVE, SUSPENDED}",
        "",
        "PRICE_ENTRY",
        "  - entry_id : int {PK}",
        "  - product_id : int {FK -> PRODUCT}",
        "  - market_id : int {FK -> MARKET}",
        "  - agent_id : int {FK -> USER}",
        "  - unit_price : decimal (KES)",
        "  - price_date : Date",
        "  - status : enum {CURRENT, HISTORICAL}",
        "  - is_anomaly : boolean",
        "  - notes : String",
        "  - UNIQUE (product_id, market_id, price_date)",
        "",
        "PRICE_ALERT",
        "  - alert_id : int {PK}",
        "  - user_id : int {FK -> USER}",
        "  - product_id : int {FK -> PRODUCT}",
        "  - market_id : int {FK -> MARKET, nullable}",
        "  - threshold_price : decimal",
        "  - alert_direction : enum {ABOVE, BELOW}",
        "  - is_active : boolean",
    ]
    diagram_block(doc, domain_model_lines)

    body_para(doc, "Key Relationships:")
    bullet_para(doc, "One REGION aggregates many MARKETs (1..*)")
    bullet_para(doc, "One PRODUCT_CATEGORY classifies many PRODUCTs (1..*)")
    bullet_para(doc, "One USER (Agent) submits many PRICE_ENTRYs (1..*)")
    bullet_para(doc, "One PRODUCT appears in many PRICE_ENTRYs and PRICE_ALERTs (1..*)")
    bullet_para(doc, "One MARKET hosts many PRICE_ENTRYs and PRICE_ALERTs (1..*)")
    bullet_para(doc, "One USER creates many PRICE_ALERTs (1..*)")

    heading2(doc, "4.6 Use Case Descriptions")

    use_cases = [
        ("Use Case 1: Login",
         "All Users",
         "User has a registered account.",
         "User submits email and password → system validates against users table → session created "
         "with userId, userRole, userName → redirect to dashboard.",
         "Invalid credentials → error message displayed, no session created."),
        ("Use Case 2: Register",
         "New User",
         "User has no existing account with the email.",
         "User submits registration form → system validates email format, password length, role "
         "selection → checks email uniqueness → inserts user record → redirects to login.",
         "Email already exists → error message displayed."),
        ("Use Case 3: Collect Market Price Data (Enter Price)",
         "Market Agent",
         "Agent is logged in. Product and market exist in the system.",
         "Agent selects product from dropdown, selects market, enters unit price and date → "
         "system validates inputs → inserts price entry → success message displayed.",
         "Duplicate entry for same product-market-date → UPSERT updates existing record."),
        ("Use Case 4: View Price List",
         "Farmer, Trader, Consumer, Agent, Admin",
         "User is logged in. At least one price entry exists.",
         "User navigates to Prices → system queries price_entries joined with products, markets, "
         "regions → displays paginated table.",
         "User applies filters (crop, market, region, date range) → filtered results displayed."),
        ("Use Case 5: Search and Filter Market Data",
         "All Registered Users",
         "User is logged in. Data exists in the system.",
         "User enters search criteria (crop name, region, date range) → system processes request "
         "using case-insensitive matching → system displays filtered results.",
         "If no results match, the system displays 'No data found.'"),
        ("Use Case 6: View Real-Time Dashboard",
         "All Registered Users",
         "User is logged in.",
         "User accesses the dashboard → system displays current prices by crop, livestock, and "
         "region → data updates dynamically.",
         "If data is unavailable, the system displays the last updated timestamp."),
        ("Use Case 7: Compare Prices",
         "Farmer, Trader, Consumer",
         "User is logged in. Same product has prices recorded at multiple markets.",
         "User selects a product and selects markets → system retrieves latest price per market "
         "→ displays comparison table.",
         "N/A"),
        ("Use Case 8: View Price Trends / Generate Charts",
         "Farmer, Trader, Consumer, Agent",
         "Historical price data exists for selected product-market pair.",
         "User selects product, market, and time window (30/90/180 days) → system queries "
         "time-series data → renders Chart.js interactive line graph.",
         "If insufficient data exists, the system displays a notification."),
        ("Use Case 9: Analyse Historical and Seasonal Trends",
         "All Registered Users",
         "Historical data is stored in the system.",
         "User selects historical analysis → system retrieves historical records → system "
         "calculates seasonal patterns → displays trend analysis results.",
         "N/A"),
        ("Use Case 10: Manage Price Alert",
         "Farmer, Trader, Consumer",
         "User is logged in.",
         "User creates alert with product, market, direction (ABOVE/BELOW), and threshold → "
         "system inserts into price_alerts → alert is evaluated on each new price entry "
         "submission; notification dispatched when triggered.",
         "User deactivates or deletes an existing alert."),
        ("Use Case 11: Generate Report",
         "Admin, Agent",
         "User is logged in with Admin or Agent role.",
         "User selects report type (Summary / CSV Export) and filters → system executes "
         "aggregate SQL query → renders report or initiates CSV file download.",
         "N/A"),
        ("Use Case 12: Aggregate Data from Multiple Sources",
         "System Administrator",
         "External APIs are configured; Admin has integration privileges.",
         "System automatically triggers scheduled data collection → retrieves, validates, and "
         "standardises data → merges data into the central database.",
         "If API connection fails, system logs the error and notifies the administrator."),
        ("Use Case 13: Manage Products / Markets / Users",
         "Admin",
         "User is logged in as Admin.",
         "Admin navigates to admin panel → views product/market/user list → adds new record "
         "or deactivates (soft-deletes) an existing one.",
         "N/A"),
    ]

    for uc_title, actor, precond, main_flow, alt_flow in use_cases:
        heading3(doc, uc_title)
        rows_uc = [
            ("Actor:", actor),
            ("Precondition:", precond),
            ("Main Flow:", main_flow),
        ]
        if alt_flow != "N/A":
            rows_uc.append(("Alternative Flow:", alt_flow))
        for label, content in rows_uc:
            p = body_para(doc, "", indent=0.25)
            r1 = p.add_run(label + " ")
            set_font(r1, bold=True)
            r2 = p.add_run(content)
            set_font(r2)

    heading2(doc, "4.7 UML Use Case Diagrams")
    body_para(doc,
        "The following diagrams use UML use case notation. These should be rendered using "
        "draw.io or StarUML as formal UML diagrams for submission. The text representations "
        "below faithfully depict the actor-system relationships."
    )

    uc_diagrams = [
        ("Use Case Diagram 1: Data Collection and Aggregation (UC3, UC12)",
         [
             "+------------------------------------------+",
             "|              AMPT System                 |",
             "|                                          |",
             "|   (UC3) Collect Market Price Data        |",
             "|   (UC12) Aggregate Data from Sources     |",
             "|                                          |",
             "+------------------------------------------+",
             "      ^                   ^",
             "      |                   |",
             "[Market Agent]    [System Administrator]",
             "                          |",
             "                  [External API/System]",
         ]),
        ("Use Case Diagram 2: Dashboard and Search (UC4, UC5, UC6)",
         [
             "+------------------------------------------+",
             "|              AMPT System                 |",
             "|                                          |",
             "|   (UC6) View Real-Time Dashboard         |",
             "|   (UC4) View Price List                  |",
             "|   (UC5) Search and Filter Market Data    |",
             "|                                          |",
             "+------------------------------------------+",
             "              ^",
             "              |",
             "     [Registered User]",
             "   (Farmer/Trader/Consumer/Agent/Admin)",
         ]),
        ("Use Case Diagram 3: Charts and Trend Analysis (UC7, UC8, UC9)",
         [
             "+------------------------------------------+",
             "|              AMPT System                 |",
             "|                                          |",
             "|   (UC8) View Price Trends / Charts       |",
             "|   (UC9) Analyse Historical Trends        |",
             "|   (UC7) Compare Prices                   |",
             "|                                          |",
             "+------------------------------------------+",
             "              ^",
             "              |",
             "     [Registered User]",
         ]),
        ("Use Case Diagram 4: Price Alerts (UC10)",
         [
             "+------------------------------------------+",
             "|              AMPT System                 |",
             "|                                          |",
             "|   (UC10) Manage Price Alert              |",
             "|          <<include>> Check Threshold     |",
             "|          <<include>> Send Notification   |",
             "|                                          |",
             "+------------------------------------------+",
             "              ^",
             "              |",
             "     [Farmer/Trader/Consumer]",
         ]),
        ("Use Case Diagram 5: Full System — All Actors and Use Cases",
         [
             "+------------------------------------------------------------------------+",
             "|                           AMPT System                                  |",
             "|                                                                        |",
             "|  (UC1)  Login              (UC3)  Collect Price Data                  |",
             "|  (UC2)  Register           (UC4)  View Price List                     |",
             "|  (UC6)  View Dashboard     (UC7)  Compare Prices                      |",
             "|  (UC5)  Search/Filter      (UC8)  View Trends/Charts                  |",
             "|  (UC9)  Seasonal Analysis  (UC10) Manage Alerts                       |",
             "|  (UC11) Generate Report    (UC13) Manage Products/Markets/Users        |",
             "|  (UC12) Aggregate Data                                                 |",
             "|                                                                        |",
             "+------------------------------------------------------------------------+",
             "       ^                    ^                      ^                ^",
             "       |                    |                      |                |",
             "[All Registered      [Market Agent]          [Admin]       [External API/",
             "    Users]         (UC3, UC4, UC8,      (UC11, UC12,        System]",
             "  (UC1-UC10)        UC11, UC13)           UC13)            (UC12)",
         ]),
    ]

    for title, lines in uc_diagrams:
        p = body_para(doc, "")
        r = p.add_run(title)
        set_font(r, bold=True, italic=True)
        diagram_block(doc, lines)

    # -----------------------------------------------------------------------
    # CHAPTER FIVE
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER FIVE: SYSTEM DESIGN SPECIFICATION")

    heading2(doc, "5.1 System Architecture")
    body_para(doc,
        "AMPT follows a three-tier web architecture implementing the Model-View-Controller "
        "(MVC) pattern, deployed on Apache Tomcat 10. The UML Component Diagram below shows "
        "the layered structure:"
    )

    arch_lines = [
        "[TIER 1: CLIENT]",
        "  Web Browser (HTML5 + CSS3 + JavaScript)",
        "  Chart.js (price visualisation)",
        "          |",
        "          | HTTP/HTTPS Requests & Responses",
        "          |",
        "[TIER 2: APPLICATION -- Apache Tomcat 10]",
        "  |",
        "  +-- CONTROLLER LAYER (Java Servlets)",
        "  |     LoginServlet | RegisterServlet | LogoutServlet",
        "  |     PriceEntryServlet | PriceListServlet | PriceEditServlet",
        "  |     PriceCompareServlet | PriceTrendServlet",
        "  |     AlertServlet | AlertCheckServlet | ReportServlet",
        "  |     ProductServlet | MarketServlet | UserMgmtServlet",
        "  |",
        "  +-- VIEW LAYER (JSP Pages + JSTL)",
        "  |     login.html | register.html | dashboard.jsp",
        "  |     prices/list.jsp | prices/entry.jsp | prices/edit.jsp",
        "  |     prices/compare.jsp | prices/trends.jsp",
        "  |     alerts/alerts.jsp | reports/summary.jsp",
        "  |     admin/*.jsp | profile/profile.jsp",
        "  |",
        "  +-- MODEL LAYER (Java Classes)",
        "        Product | Market | Region | Category",
        "        PriceEntry | PriceAlert | User",
        "        DBConnection (Utility) | AlgorithmUtils (Utility)",
        "          |",
        "          | JDBC (PostgreSQL Driver 42.7.1)",
        "          |",
        "[TIER 3: DATABASE -- PostgreSQL 16]",
        "  regions | markets | product_categories | products",
        "  users | price_entries | price_alerts | seasons",
        "  audit_logs | notifications | reports | market_agents",
    ]
    diagram_block(doc, arch_lines)

    body_para(doc, "Request-Response Cycle:")
    numbered_para(doc, "Browser sends HTTP GET/POST to a URL mapped to a Servlet.")
    numbered_para(doc, "Servlet authenticates the session (HttpSession check).")
    numbered_para(doc, "Servlet executes JDBC queries using PreparedStatement.")
    numbered_para(doc, "Servlet sets result data as request attributes.")
    numbered_para(doc, "Servlet forwards request to the corresponding JSP.")
    numbered_para(doc, "JSP renders HTML using JSTL tags.")
    numbered_para(doc, "Rendered HTML returned to browser.")

    heading2(doc, "5.2 Database Design")

    heading3(doc, "5.2.1 Entity-Relationship Overview")
    body_para(doc, "The AMPT database consists of 12 tables in PostgreSQL 16:")

    p = body_para(doc, "")
    r = p.add_run("Reference / Lookup Tables:")
    set_font(r, bold=True)
    bullet_para(doc, "regions — Kenya's 8 geographic regions")
    bullet_para(doc, "product_categories — 9 commodity classification categories")
    bullet_para(doc, "seasons — agricultural season definitions")

    p = body_para(doc, "")
    r = p.add_run("Core Operational Tables:")
    set_font(r, bold=True)
    bullet_para(doc, "products → depends on product_categories (35 commodities)")
    bullet_para(doc, "markets → depends on regions (29 markets)")
    bullet_para(doc, "users → independent (authentication anchor)")

    p = body_para(doc, "")
    r = p.add_run("Transactional Tables:")
    set_font(r, bold=True)
    bullet_para(doc, "price_entries → depends on products, markets, users (agent), seasons")
    bullet_para(doc, "price_alerts → depends on users, products, markets")
    bullet_para(doc, "market_agents → depends on users")

    p = body_para(doc, "")
    r = p.add_run("Audit / Notification Tables:")
    set_font(r, bold=True)
    bullet_para(doc, "audit_logs → depends on users")
    bullet_para(doc, "notifications → depends on users, price_alerts")
    bullet_para(doc, "reports → depends on users")

    heading3(doc, "5.2.2 Key Table Definitions (SQL DDL)")

    sql_products = [
        "CREATE TABLE products (",
        "    product_id     SERIAL PRIMARY KEY,",
        "    product_name   VARCHAR(255) NOT NULL,",
        "    local_name     VARCHAR(255),",
        "    standard_unit  VARCHAR(50) NOT NULL,",
        "    category_id    INTEGER REFERENCES product_categories(category_id),",
        "    is_active      BOOLEAN DEFAULT TRUE",
        ");",
    ]
    code_block(doc, sql_products)
    doc.add_paragraph()

    sql_markets = [
        "CREATE TABLE markets (",
        "    market_id       SERIAL PRIMARY KEY,",
        "    market_name     VARCHAR(255) NOT NULL,",
        "    town            VARCHAR(100),",
        "    latitude        NUMERIC(9,6),",
        "    longitude       NUMERIC(9,6),",
        "    operating_days  VARCHAR(100),",
        "    operating_hours VARCHAR(100),",
        "    status          VARCHAR(20) DEFAULT 'ACTIVE',",
        "    region_id       INTEGER REFERENCES regions(region_id)",
        ");",
    ]
    code_block(doc, sql_markets)
    doc.add_paragraph()

    sql_prices = [
        "CREATE TABLE price_entries (",
        "    entry_id             SERIAL PRIMARY KEY,",
        "    product_id           INTEGER REFERENCES products(product_id),",
        "    market_id            INTEGER REFERENCES markets(market_id),",
        "    agent_id             INTEGER REFERENCES market_agents(agent_id),",
        "    unit_price           NUMERIC(12,2) NOT NULL,",
        "    currency             VARCHAR(3) DEFAULT 'KES',",
        "    price_date           DATE NOT NULL,",
        "    status               VARCHAR(20) DEFAULT 'CURRENT',",
        "    is_anomaly           BOOLEAN DEFAULT FALSE,",
        "    notes                TEXT,",
        "    submission_timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,",
        "    UNIQUE (product_id, market_id, price_date)",
        ");",
    ]
    code_block(doc, sql_prices)
    doc.add_paragraph()

    sql_alerts = [
        "CREATE TABLE price_alerts (",
        "    alert_id        SERIAL PRIMARY KEY,",
        "    user_id         INTEGER REFERENCES users(user_id),",
        "    product_id      INTEGER REFERENCES products(product_id),",
        "    market_id       INTEGER REFERENCES markets(market_id),",
        "    threshold_price NUMERIC(12,2) NOT NULL,",
        "    alert_direction VARCHAR(10),",
        "    is_active       BOOLEAN DEFAULT TRUE,",
        "    created_date    TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
        ");",
    ]
    code_block(doc, sql_alerts)
    doc.add_paragraph()

    sql_index = [
        "-- Performance Index",
        "CREATE INDEX idx_price_lookup",
        "  ON price_entries (product_id, market_id, price_date);",
    ]
    code_block(doc, sql_index)
    doc.add_paragraph()

    body_para(doc,
        "This composite index accelerates the most frequent query patterns: filtering by "
        "product, filtering by market, and ordering by date for trend analysis."
    )

    heading3(doc, "5.2.3 Seed Data Summary")

    seed_rows = [
        ("Table", "Records"),
        ("regions", "8 (all Kenyan regions)"),
        ("product_categories", "9"),
        ("products", "35"),
        ("markets", "29 (major markets per region)"),
    ]
    tbl_seed = doc.add_table(rows=len(seed_rows), cols=2)
    tbl_seed.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, rd in enumerate(seed_rows):
        row = tbl_seed.rows[i]
        for j, val in enumerate(rd):
            row.cells[j].width = Inches(2.0)
            table_cell(row.cells[j], val, bold=(i == 0), size=11)
    style_table(tbl_seed)

    heading2(doc, "5.3 Flowcharts and Activity Diagrams")
    body_para(doc,
        "The following diagrams use UML Activity Diagram notation. Render using draw.io or "
        "StarUML for formal submission."
    )

    heading3(doc, "5.3.1 Activity Diagram: Price Entry Workflow")
    price_entry_flow = [
        "[Start]",
        "  |",
        "  v",
        "[Market Agent logs in]",
        "  |",
        "  v",
        "[Navigate to /prices/entry]",
        "  |",
        "  v",
        "[GET: Servlet loads products and markets from DB]",
        "  |",
        "  v",
        "[JSP renders entry form with populated dropdowns]",
        "  |",
        "  v",
        "[Agent selects product, market, enters price & date]",
        "  |",
        "  v",
        "[POST submitted]",
        "  |",
        "  v",
        "<Validate inputs: price > 0, date not null>",
        "  |          |",
        " [Yes]      [No] --> [Set error attribute, re-render form]",
        "  |",
        "  v",
        "[INSERT INTO price_entries ON CONFLICT DO UPDATE]",
        "  |          |",
        "[Success]  [Fail] --> [Set DB error, redirect to entry form]",
        "  |",
        "  v",
        "[Redirect to /prices/list with success message]",
        "  |",
        "  v",
        "[End]",
    ]
    diagram_block(doc, price_entry_flow)

    heading3(doc, "5.3.2 Activity Diagram: User Authentication Workflow")
    auth_flow = [
        "[User accesses any protected URL]",
        "  |",
        "  v",
        "<Session has userId?>",
        "  |              |",
        " [Yes]          [No]",
        "  |              v",
        "  |         [Redirect to /login]",
        "  |              |",
        "  |         [User submits email + password]",
        "  |              |",
        "  |         <SELECT user WHERE email=? AND password_hash=bcrypt(?)>",
        "  |              |",
        "  |         <User found AND status=ACTIVE?>",
        "  |              |              |",
        "  |            [Yes]           [No]",
        "  |              |              v",
        "  |         [Set session]  [Set error: Invalid credentials]",
        "  |         [Redirect to   [Re-render login page]",
        "  |          dashboard]",
        "  v",
        "[Servlet processes request normally]",
        "  |",
        "  v",
        "[End]",
    ]
    diagram_block(doc, auth_flow)

    heading3(doc, "5.3.3 Flowchart: Price Alert Check Workflow")
    alert_flow = [
        "[Start]",
        "  |",
        "  v",
        "[SELECT all active price_alerts]",
        "  |",
        "  v",
        "[FOR EACH alert:]",
        "  |",
        "  v",
        "[SELECT latest price for alert.product_id, alert.market_id]",
        "  |",
        "  v",
        "<direction = 'ABOVE' AND current_price > threshold?>",
        "  |                    |",
        " [Yes]               [No]",
        "  |                    v",
        "  |         <direction = 'BELOW' AND current_price < threshold?>",
        "  |                    |              |",
        "  |                  [Yes]           [No]",
        "  |                    |              v",
        "  |                    |      [No action -- condition not met]",
        "  |                    |",
        "  v                    v",
        "[Trigger notification: insert into notifications table]",
        "  |",
        "  v",
        "[Next alert]",
        "  |",
        "  v",
        "[End]",
    ]
    diagram_block(doc, alert_flow)

    heading2(doc, "5.4 Interface Design Overview")
    body_para(doc, "The system interface follows a consistent design language:")
    bullet_para(doc, "Header/Navbar: Persistent navigation bar with brand logo, system name, and role-aware navigation links.")
    bullet_para(doc, "Data Tables: Semantic table elements with product name, unit price (KES), market, region, date, and action buttons.")
    bullet_para(doc, "Forms: Labelled inputs with placeholder text, submit/cancel button pairs.")
    bullet_para(doc, "Charts: Chart.js line and bar charts for price trends with 30/90/180-day selectors.")
    bullet_para(doc, "Colour Palette: Dark green primary (#1a4731), warm off-white background — aligned with agricultural branding.")
    bullet_para(doc, "Typography: Playfair Display (headings), IBM Plex Mono (prices/data), Source Serif 4 (body).")
    bullet_para(doc, "Responsiveness: Flexbox layout with mobile breakpoints for navbar collapse.")

    # -----------------------------------------------------------------------
    # CHAPTER SIX
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER SIX: SYSTEM IMPLEMENTATION AND TESTING")

    heading2(doc, "6.1 Implementation Language")
    body_para(doc,
        "AMPT is implemented in Java SE 17 for the backend, with HTML5, CSS3, and vanilla "
        "JavaScript for the frontend. Java was selected for the following reasons:"
    )
    bullet_para(doc,
        "Platform Maturity: Java has decades of enterprise web application history with "
        "Jakarta EE, providing well-specified APIs for servlet, JSP, and JDBC programming."
    )
    bullet_para(doc,
        "Type Safety: Java's static typing reduces runtime errors in data handling code, "
        "particularly important for financial data (prices) where type errors produce "
        "incorrect results."
    )
    bullet_para(doc,
        "Academic Alignment: The course requires Java implementation. Java Servlets demonstrate "
        "fundamental HTTP request-response handling, session management, and server-side "
        "rendering without abstraction layers obscuring the underlying mechanics."
    )
    bullet_para(doc,
        "JDBC: Direct database access with PreparedStatement provides transparent SQL "
        "execution, making query optimisation and security measures explicit and auditable."
    )

    heading2(doc, "6.2 Key Language Features Used")

    feat_rows = [
        ("Feature", "Usage in AMPT"),
        ("Generics (List<String[]>)", "Type-safe collections for price list rows returned from JDBC ResultSet"),
        ("Try-with-resources", "Automatic Connection, PreparedStatement, ResultSet closing — prevents resource leaks"),
        ("HttpSession", "Session management for user authentication state across requests"),
        ("PreparedStatement", "Parameterised SQL queries preventing SQL injection"),
        ("RequestDispatcher.forward()", "MVC forwarding from Servlet controller to JSP view"),
        ("@WebServlet annotation", "Declarative URL pattern mapping without XML verbosity"),
        ("BigDecimal", "Precise arithmetic for financial price values, avoiding floating-point errors"),
        ("JSONB via JDBC", "Flexible audit log change_details stored as JSON in PostgreSQL"),
        ("response.setContentType(\"text/csv\")", "Streaming CSV report download with appropriate MIME type"),
    ]
    tbl_feat = doc.add_table(rows=len(feat_rows), cols=2)
    tbl_feat.alignment = WD_TABLE_ALIGNMENT.LEFT
    fw = [Inches(2.5), Inches(4.0)]
    for i, rd in enumerate(feat_rows):
        row = tbl_feat.rows[i]
        for j, val in enumerate(rd):
            row.cells[j].width = fw[j]
            table_cell(row.cells[j], val, bold=(i == 0), size=10)
    style_table(tbl_feat)

    heading2(doc, "6.3 Search Algorithms Implemented")
    body_para(doc,
        "AMPT includes an AlgorithmUtils class implementing standard algorithms in Java, "
        "demonstrating core computer science concepts within the application domain."
    )

    heading3(doc, "6.3.1 Linear Search — O(n)")
    code_block(doc, [
        "public static int linearSearch(List<String> list, String target) {",
        "    for (int i = 0; i < list.size(); i++) {",
        "        if (list.get(i).equalsIgnoreCase(target)) {",
        "            return i;",
        "        }",
        "    }",
        "    return -1;",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Searching unsorted product name lists for a matching entry.")
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Complexity: ")
    set_font(r, bold=True)
    r2 = p.add_run("Time O(n), Space O(1).")
    set_font(r2)

    heading3(doc, "6.3.2 Binary Search — O(log n)")
    code_block(doc, [
        "public static int binarySearch(List<String> sortedList, String target) {",
        "    int low = 0, high = sortedList.size() - 1;",
        "    while (low <= high) {",
        "        int mid = (low + high) / 2;",
        "        int cmp = sortedList.get(mid).compareToIgnoreCase(target);",
        "        if (cmp == 0) return mid;",
        "        else if (cmp < 0) low = mid + 1;",
        "        else high = mid - 1;",
        "    }",
        "    return -1;",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Searching sorted market name lists and historical price date arrays.")
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Complexity: ")
    set_font(r, bold=True)
    r2 = p.add_run("Time O(log n), Space O(1). Prerequisite: sorted input list.")
    set_font(r2)

    heading3(doc, "6.3.3 Bubble Sort — O(n²)")
    code_block(doc, [
        "public static void bubbleSort(List<Double> prices) {",
        "    int n = prices.size();",
        "    for (int i = 0; i < n - 1; i++) {",
        "        boolean swapped = false;",
        "        for (int j = 0; j < n - i - 1; j++) {",
        "            if (prices.get(j) > prices.get(j + 1)) {",
        "                double tmp = prices.get(j);",
        "                prices.set(j, prices.get(j + 1));",
        "                prices.set(j + 1, tmp);",
        "                swapped = true;",
        "            }",
        "        }",
        "        if (!swapped) break;",
        "    }",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Sorting small price arrays for anomaly detection.")
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Complexity: ")
    set_font(r, bold=True)
    r2 = p.add_run("Time O(n²) worst-case, O(n) best-case. Space O(1).")
    set_font(r2)

    heading3(doc, "6.3.4 Quick Sort — O(n log n) average")
    code_block(doc, [
        "public static void quickSort(List<Double> prices, int low, int high) {",
        "    if (low < high) {",
        "        int pi = partition(prices, low, high);",
        "        quickSort(prices, low, pi - 1);",
        "        quickSort(prices, pi + 1, high);",
        "    }",
        "}",
        "",
        "private static int partition(List<Double> prices, int low, int high) {",
        "    double pivot = prices.get(high);",
        "    int i = low - 1;",
        "    for (int j = low; j < high; j++) {",
        "        if (prices.get(j) <= pivot) {",
        "            i++;",
        "            double tmp = prices.get(i);",
        "            prices.set(i, prices.get(j));",
        "            prices.set(j, tmp);",
        "        }",
        "    }",
        "    double tmp = prices.get(i + 1);",
        "    prices.set(i + 1, prices.get(high));",
        "    prices.set(high, tmp);",
        "    return i + 1;",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Sorting larger price datasets for trend analysis and median calculation.")
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Complexity: ")
    set_font(r, bold=True)
    r2 = p.add_run("Time O(n log n) average, O(n²) worst-case. Space O(log n).")
    set_font(r2)

    heading3(doc, "6.3.5 Min/Max Finder — O(n)")
    code_block(doc, [
        "public static double[] findMinMaxPrices(List<Double> prices) {",
        "    double min = Double.MAX_VALUE, max = Double.MIN_VALUE;",
        "    for (double p : prices) {",
        "        if (p < min) min = p;",
        "        if (p > max) max = p;",
        "    }",
        "    return new double[]{min, max};",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Computing price range for report summary statistics.")
    set_font(r2)
    p = body_para(doc, "")
    r = p.add_run("Complexity: ")
    set_font(r, bold=True)
    r2 = p.add_run("Time O(n), Space O(1).")
    set_font(r2)

    heading3(doc, "6.3.6 Median Price Calculation")
    code_block(doc, [
        "public static double getMedianPrice(List<Double> prices) {",
        "    List<Double> sorted = new ArrayList<>(prices);",
        "    quickSort(sorted, 0, sorted.size() - 1);",
        "    int n = sorted.size();",
        "    if (n % 2 == 1) return sorted.get(n / 2);",
        "    else return (sorted.get((n - 1) / 2) + sorted.get(n / 2)) / 2.0;",
        "}",
    ])
    p = body_para(doc, "")
    r = p.add_run("Use: ")
    set_font(r, bold=True)
    r2 = p.add_run("Dashboard display of median market price for a commodity.")
    set_font(r2)

    heading3(doc, "6.3.7 Database-Level Algorithms")
    body_para(doc,
        "Beyond in-memory algorithms, AMPT leverages PostgreSQL's query planner:"
    )
    bullet_para(doc, "ORDER BY price_date DESC — B-tree index scan for trend data retrieval.")
    bullet_para(doc, "ILIKE '%pattern%' — case-insensitive pattern matching for search filters.")
    bullet_para(doc, "GROUP BY product_id, market_id with AVG(), MIN(), MAX(), COUNT() — aggregate functions for report generation.")
    bullet_para(doc, "LIMIT 1 with ORDER BY price_date DESC — latest price retrieval using the composite index idx_price_lookup.")
    bullet_para(doc, "WHERE price_date >= CURRENT_DATE - INTERVAL 'N days' — time-window filtering for trend queries.")

    heading2(doc, "6.4 Testing Strategy")

    heading3(doc, "6.4.1 Unit Testing")
    body_para(doc,
        "Individual utility functions (AlgorithmUtils) tested in isolation using JUnit:"
    )
    bullet_para(doc, "AlgorithmUtils.linearSearch() — tested with found, not-found, and case-insensitive cases.")
    bullet_para(doc, "AlgorithmUtils.binarySearch() — tested with sorted lists, empty lists, single element.")
    bullet_para(doc, "AlgorithmUtils.bubbleSort() — verified against pre-sorted, reverse-sorted, and random arrays.")
    bullet_para(doc, "AlgorithmUtils.quickSort() — verified sort correctness and stability.")

    heading3(doc, "6.4.2 Integration Testing")
    body_para(doc,
        "End-to-end workflows tested by running the system on Tomcat against a test "
        "PostgreSQL database:"
    )
    bullet_para(doc, "Price entry → confirmation → appears in price list.")
    bullet_para(doc, "Price filter by crop name → correct filtered results.")
    bullet_para(doc, "Alert creation → alert triggered when price inserted above threshold.")
    bullet_para(doc, "Admin product add → product appears in price entry dropdown.")
    bullet_para(doc, "CSV report download → valid CSV file with correct headers and data.")

    heading3(doc, "6.4.3 Security Testing")
    bullet_para(doc,
        "SQL injection attempt: entering '; DROP TABLE price_entries; -- into filter "
        "fields — verified safe due to PreparedStatement binding."
    )
    bullet_para(doc,
        "Session bypass attempt: accessing /admin/users without login → verified redirect "
        "to /login."
    )
    bullet_para(doc,
        "Role bypass attempt: logged in as FARMER, directly accessing /prices/entry via URL "
        "→ verified servlet returns 403 based on role check."
    )

    heading3(doc, "6.4.4 User Acceptance Testing (UAT)")
    body_para(doc, "Test scenarios covering:")
    bullet_para(doc, "Login / logout flow.")
    bullet_para(doc, "New user registration with valid and invalid inputs.")
    bullet_para(doc, "Price entry, editing, and deletion.")
    bullet_para(doc, "Price comparison and trend chart rendering.")
    bullet_para(doc, "Alert creation and notification.")
    bullet_para(doc, "Report generation and CSV download.")
    bullet_para(doc, "Admin CRUD for products, markets, and users.")

    heading2(doc, "6.5 Source Code")
    body_para(doc,
        "The full system source code is submitted on a CD attached to this report. "
        "The CD contains:"
    )

    cd_contents = [
        "CD_CONTENTS/",
        "├── src/                    (Java Servlet and Model source files)",
        "├── WebContent/             (JSP views, CSS, JavaScript)",
        "│   ├── WEB-INF/web.xml     (Servlet URL mappings)",
        "│   └── css/style.css       (Global stylesheet)",
        "├── db/",
        "│   ├── database_setup_fixed.sql   (Schema creation script)",
        "│   └── seed_data.sql              (Reference data: regions, products, markets)",
        "├── pom.xml                 (Maven build configuration)",
        "└── QUICKSTART.md           (Setup and deployment instructions)",
    ]
    code_block(doc, cd_contents)
    doc.add_paragraph()

    body_para(doc, "To run the system from the CD:")
    numbered_para(doc, "Install Java 17, PostgreSQL 16, Apache Maven 3.8+, and Apache Tomcat 10.")
    numbered_para(doc, "Create database: createdb agri_price_tracker")
    numbered_para(doc, "Run schema: psql -U postgres -d agri_price_tracker -f db/database_setup_fixed.sql")
    numbered_para(doc, "Run seed data: psql -U postgres -d agri_price_tracker -f db/seed_data.sql")
    numbered_para(doc, "Build: mvn clean package")
    numbered_para(doc, "Deploy: Copy target/agri-price-tracker.war to Tomcat webapps/ directory.")
    numbered_para(doc, "Access: http://localhost:8080/agri-price-tracker")

    # -----------------------------------------------------------------------
    # CHAPTER SEVEN
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "CHAPTER SEVEN: CONCLUSION AND RECOMMENDATION")

    heading2(doc, "7.1 Conclusion")
    body_para(doc,
        "The Agricultural Market Price Tracker (AMPT) — AgriPrice KE — has been designed and "
        "implemented as a full-featured, role-based web application addressing the critical "
        "problem of agricultural market price information asymmetry in Kenya."
    )
    body_para(doc, "The system achieves all stated specific objectives:")
    numbered_para(doc,
        "Requirements Analysis: A comprehensive set of 28 functional requirements and 13 "
        "non-functional requirements was elicited through stakeholder analysis, literature "
        "review, and use case modelling using Object-Oriented Analysis techniques."
    )
    numbered_para(doc,
        "Database Design: A normalised relational schema of 12 tables was designed in "
        "PostgreSQL 16, covering the full domain — geographic regions, markets, commodity "
        "categories, products, user accounts, price entries, alerts, audit logs, and reports "
        "— seeded with data for 8 Kenyan regions, 35 agricultural commodities, and 29 major "
        "markets."
    )
    numbered_para(doc,
        "Implementation: A role-based web application was implemented using Java 17 Servlets "
        "+ JSP (MVC), with 14 servlet controllers, 14 JSP views, and 6 model classes. "
        "Role-based access control, session management, and SQL injection prevention are "
        "fully implemented."
    )
    numbered_para(doc,
        "Analytical Features: Price comparison across markets, time-series trend visualisation "
        "using Chart.js, configurable price alerts, and CSV report export were successfully "
        "implemented and tested."
    )
    numbered_para(doc,
        "Algorithm Implementation: Linear search, binary search, bubble sort (with early-exit "
        "optimisation), quicksort, median calculation, and min/max finding algorithms were "
        "implemented in Java within the AlgorithmUtils class, complementing database-level "
        "sorting and filtering."
    )
    numbered_para(doc,
        "Testing: The system was validated through unit, integration, security, and user "
        "acceptance testing, with all critical functional requirements verified."
    )
    body_para(doc,
        "The project demonstrates practical competence in software engineering: version control "
        "with Git, collaborative development, iterative development with code reviews, and "
        "structured documentation aligned with professional standards."
    )

    heading2(doc, "7.2 Recommendations")
    body_para(doc,
        "Based on the development experience and testing outcomes, the following enhancements "
        "are recommended for future development and deployment of AMPT:"
    )
    numbered_para(doc,
        "SMS Alert Integration: Integration with Africa's Talking SMS API would deliver "
        "immediate value to farmers who may not check the web interface daily. This represents "
        "the highest-impact single enhancement."
    )
    numbered_para(doc,
        "Mobile Application: A Progressive Web App (PWA) wrapper or native Android application "
        "would extend reach to market agents operating in areas with limited laptop/desktop "
        "access."
    )
    numbered_para(doc,
        "Automated Data Collection: Partnering with county government agriculture departments "
        "to automate price data submission via mobile data collection forms or APIs would "
        "improve data freshness."
    )
    numbered_para(doc,
        "Anomaly Detection Enhancement: Implementing a statistical anomaly detection algorithm "
        "(Z-score or IQR-based outlier detection) that automatically flags prices deviating "
        "significantly from the 30-day mean would improve data quality."
    )
    numbered_para(doc,
        "Multi-Language Support: Adding Kiswahili localisation would significantly expand "
        "accessibility for the target demographic, as Kenya's agricultural population is most "
        "comfortable in Kiswahili."
    )
    numbered_para(doc,
        "Connection Pooling: Replacing DriverManager with a connection pool (HikariCP or "
        "Apache DBCP) would substantially improve performance and resilience under concurrent "
        "production load."
    )
    numbered_para(doc,
        "Cloud Deployment: Deploying AMPT to a cloud platform (AWS, Google Cloud, or Safaricom "
        "Cloud) would make the system accessible nationwide, aligning with Kenya's Digital "
        "Economy Blueprint objectives."
    )

    # -----------------------------------------------------------------------
    # REFERENCES
    # -----------------------------------------------------------------------
    doc.add_page_break()
    heading1(doc, "REFERENCES")

    references = [
        "Chuka University (2025). Course Outline: Software Engineering / System Development. "
        "Department of Computer Science, School of Computing and Informatics.",

        "Communications Authority of Kenya (2024). Sector Statistics Report Q2 2023/2024. "
        "Nairobi: CA Kenya.",

        "Eriksson, H. E., & Penker, M. (1998). UML Toolkit. New York: John Wiley & Sons.",

        "Food and Agriculture Organization of the United Nations (2021). Market Information "
        "Systems and Agricultural Commodities Markets. FAO Agricultural Development Economics "
        "Working Paper 21-07. Rome: FAO.",

        "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements "
        "of Reusable Object-Oriented Software. Reading, MA: Addison-Wesley.",

        "Government of Kenya (2021). The Big Four Agenda: Overview. Nairobi: GoK.",

        "Horstmann, C. S. (2019). Core Java, Volume II — Advanced Features (11th ed.). "
        "Hoboken, NJ: Pearson Education.",

        "Jensen, K. (2022). Agricultural Market Information Systems in Sub-Saharan Africa: "
        "A Systematic Review. Journal of Agricultural Informatics, 13(2), 1–18.",

        "Kenya Ministry of Agriculture, Livestock and Fisheries (2022). Agricultural Sector "
        "Transformation and Growth Strategy (ASTGS) 2021–2026. Nairobi: MOALF.",

        "Kenya National Bureau of Statistics (2023). Kenya Statistical Abstract 2023. "
        "Nairobi: KNBS.",

        "Larman, C. (2004). Applying UML and Patterns: An Introduction to Object-Oriented "
        "Analysis and Design (3rd ed.). Upper Saddle River, NJ: Prentice Hall.",

        "Mwebaze, P., & Okello, J. (2020). Price Information Systems and Smallholder Market "
        "Participation in East Africa. African Journal of Agricultural and Resource Economics, "
        "15(3), 210–228.",

        "Oracle Corporation (2024). Java SE 17 Documentation. Oracle.",

        "Ochieng, E. et al. (2020). Mobile-Based Price Notification System for Maize Farmers "
        "in Western Kenya. African Journal of Computing & ICT, 13(1), 45–56.",

        "PostgreSQL Global Development Group (2024). PostgreSQL 16 Documentation.",

        "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's "
        "Approach (9th ed.). New York: McGraw-Hill Education.",

        "Rumbaugh, J., Jacobson, I., & Booch, G. (2004). The Unified Modeling Language "
        "Reference Manual (2nd ed.). Boston: Addison-Wesley.",

        "Sommerville, I. (2016). Software Engineering (10th ed.). Harlow: Pearson Education.",

        "World Bank (2022). Agriculture and Food: Kenya Overview. Washington, DC: The World "
        "Bank Group.",
    ]

    for ref in references:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(ref)
        set_font(run)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    build_document()
